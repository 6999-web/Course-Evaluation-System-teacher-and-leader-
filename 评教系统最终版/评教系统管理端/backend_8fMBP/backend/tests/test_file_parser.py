"""
文件解析器单元测试

测试 FileParser 类的各种功能
"""

import os
import sys
import pytest
from pathlib import Path

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent.parent / "app"))

from utils.file_parser import (
    FileParser,
    UnsupportedFormatError,
    EmptyFileError,
    FileCorruptedError,
    parse_file,
    get_file_info,
    is_supported_format
)


class TestFileParser:
    """文件解析器测试类"""
    
    @pytest.fixture
    def parser(self):
        """创建文件解析器实例"""
        return FileParser()
    
    @pytest.fixture
    def test_files_dir(self):
        """测试文件目录"""
        return Path(__file__).parent / "test_files"
    
    def test_parser_initialization(self, parser):
        """测试解析器初始化"""
        assert parser is not None
        assert hasattr(parser, 'parse_file')
        assert hasattr(parser, 'parse_docx')
        assert hasattr(parser, 'parse_pdf')
        assert hasattr(parser, 'parse_pptx')
    
    def test_supported_formats(self, parser):
        """测试支持的文件格式"""
        assert 'docx' in parser.SUPPORTED_FORMATS
        assert 'pdf' in parser.SUPPORTED_FORMATS
        assert 'pptx' in parser.SUPPORTED_FORMATS
        assert 'ppt' in parser.SUPPORTED_FORMATS
    
    def test_is_supported_format(self):
        """测试格式支持检查"""
        assert is_supported_format('docx') is True
        assert is_supported_format('pdf') is True
        assert is_supported_format('pptx') is True
        assert is_supported_format('txt') is False
        assert is_supported_format('jpg') is False
    
    def test_unsupported_format_error(self, parser, tmp_path):
        """测试不支持的文件格式错误处理"""
        # 创建一个不支持的文件类型
        test_file = tmp_path / "test.txt"
        test_file.write_text("This is a text file")
        
        with pytest.raises(UnsupportedFormatError) as exc_info:
            parser.parse_file(str(test_file))
        
        assert "不支持的文件格式" in str(exc_info.value)
    
    def test_file_not_found_error(self, parser):
        """测试文件不存在错误"""
        with pytest.raises(FileNotFoundError):
            parser.parse_file("/nonexistent/file.docx")
    
    def test_empty_file_error(self, parser, tmp_path):
        """测试空文件错误处理"""
        # 创建一个空的 DOCX 文件
        empty_file = tmp_path / "empty.docx"
        empty_file.touch()
        
        with pytest.raises(EmptyFileError):
            parser.parse_file(str(empty_file))
    
    def test_parse_docx_basic(self, parser, tmp_path):
        """测试基本的 DOCX 文件解析"""
        # 创建一个简单的 DOCX 文件
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("这是一个测试文档")
            doc.add_paragraph("包含多个段落")
            doc.add_heading("标题", level=1)
            doc.add_paragraph("更多内容")
            
            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))
            
            # 解析文件
            content = parser.parse_docx(str(test_file))
            
            assert "这是一个测试文档" in content
            assert "包含多个段落" in content
            assert "标题" in content
            assert "更多内容" in content
            
        except ImportError:
            pytest.skip("python-docx 未安装")
    
    def test_parse_docx_with_table(self, parser, tmp_path):
        """测试包含表格的 DOCX 文件解析"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("文档内容")
            
            # 添加表格
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "单元格1"
            table.cell(0, 1).text = "单元格2"
            table.cell(1, 0).text = "单元格3"
            table.cell(1, 1).text = "单元格4"
            
            test_file = tmp_path / "test_table.docx"
            doc.save(str(test_file))
            
            # 解析文件
            content = parser.parse_docx(str(test_file))
            
            assert "文档内容" in content
            assert "单元格1" in content
            assert "单元格2" in content
            assert "单元格3" in content
            assert "单元格4" in content
            
        except ImportError:
            pytest.skip("python-docx 未安装")
    
    def test_get_file_info(self, parser, tmp_path):
        """测试获取文件信息"""
        # 创建测试文件
        test_file = tmp_path / "test.docx"
        test_file.write_text("test content")
        
        info = parser.get_file_info(str(test_file))
        
        assert info['file_name'] == "test.docx"
        assert info['file_type'] == "docx"
        assert info['file_size'] > 0
        assert info['is_supported'] is True
        assert 'file_size_mb' in info
        assert 'modified_time' in info
    
    def test_get_file_info_unsupported(self, parser, tmp_path):
        """测试获取不支持格式文件的信息"""
        test_file = tmp_path / "test.txt"
        test_file.write_text("test content")
        
        info = parser.get_file_info(str(test_file))
        
        assert info['file_type'] == "txt"
        assert info['is_supported'] is False
        assert info['format_name'] == "未知格式"
    
    def test_convenience_functions(self, tmp_path):
        """测试便捷函数"""
        try:
            from docx import Document
            
            # 创建测试文件
            doc = Document()
            doc.add_paragraph("测试内容")
            test_file = tmp_path / "test.docx"
            doc.save(str(test_file))
            
            # 测试 parse_file 便捷函数
            content = parse_file(str(test_file))
            assert "测试内容" in content
            
            # 测试 get_file_info 便捷函数
            info = get_file_info(str(test_file))
            assert info['file_name'] == "test.docx"
            
        except ImportError:
            pytest.skip("python-docx 未安装")
    
    def test_parse_file_with_explicit_type(self, parser, tmp_path):
        """测试显式指定文件类型的解析"""
        try:
            from docx import Document
            
            # 创建一个没有扩展名的文件
            doc = Document()
            doc.add_paragraph("测试内容")
            test_file = tmp_path / "test_no_ext"
            doc.save(str(test_file))
            
            # 显式指定文件类型
            content = parser.parse_file(str(test_file), file_type='docx')
            assert "测试内容" in content
            
        except ImportError:
            pytest.skip("python-docx 未安装")


class TestFileParserEdgeCases:
    """文件解析器边界情况测试"""
    
    @pytest.fixture
    def parser(self):
        return FileParser()
    
    def test_parse_docx_empty_content(self, parser, tmp_path):
        """测试解析只有空段落的 DOCX 文件"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("")  # 空段落
            doc.add_paragraph("   ")  # 只有空格的段落
            
            test_file = tmp_path / "empty_content.docx"
            doc.save(str(test_file))
            
            # 尝试解析，可能会返回空内容或抛出异常
            try:
                content = parser.parse_docx(str(test_file))
                # 如果没有抛出异常，内容应该为空或只有空白
                assert content.strip() == "" or len(content.strip()) < 10
            except EmptyFileError:
                # 抛出 EmptyFileError 也是可以接受的
                pass
                
        except ImportError:
            pytest.skip("python-docx 未安装")
    
    def test_parse_large_file(self, parser, tmp_path):
        """测试解析大文件"""
        try:
            from docx import Document
            
            doc = Document()
            # 添加大量内容
            for i in range(100):
                doc.add_paragraph(f"这是第 {i+1} 段内容，包含一些测试文本。" * 10)
            
            test_file = tmp_path / "large.docx"
            doc.save(str(test_file))
            
            # 解析文件
            content = parser.parse_docx(str(test_file))
            
            assert len(content) > 1000
            assert "这是第 1 段内容" in content
            assert "这是第 100 段内容" in content
            
        except ImportError:
            pytest.skip("python-docx 未安装")
    
    def test_parse_file_with_special_characters(self, parser, tmp_path):
        """测试解析包含特殊字符的文件"""
        try:
            from docx import Document
            
            doc = Document()
            doc.add_paragraph("特殊字符测试：@#$%^&*()")
            doc.add_paragraph("中文字符：你好世界")
            doc.add_paragraph("数字：1234567890")
            doc.add_paragraph("符号：！@#￥%……&*（）")
            
            test_file = tmp_path / "special_chars.docx"
            doc.save(str(test_file))
            
            content = parser.parse_docx(str(test_file))
            
            assert "@#$%^&*()" in content
            assert "你好世界" in content
            assert "1234567890" in content
            
        except ImportError:
            pytest.skip("python-docx 未安装")


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
